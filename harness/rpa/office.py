"""
Office automation integrations: Outlook (win32com), Word (python-docx), PDF (pypdf).
Platform-specific dependencies loaded lazily.
"""

from pathlib import Path
from typing import Any, Dict, List, Optional

from harness.logger import HarnessLogger


class OutlookHandler:
    def __init__(self, logger: Optional[HarnessLogger] = None):
        self.logger = logger or HarnessLogger("outlook")
        self._outlook = None
        self._namespace = None

    def connect(self):
        try:
            import win32com.client
            self._outlook = win32com.client.Dispatch("Outlook.Application")
            self._namespace = self._outlook.GetNamespace("MAPI")
            self.logger.info("Connected to Outlook")
        except ImportError:
            self.logger.error("pywin32 not installed. Outlook integration requires Windows + pywin32.")
            raise
        except Exception as e:
            self.logger.error(f"Failed to connect to Outlook: {e}")
            raise

    def search_emails(
        self,
        subject_contains: Optional[str] = None,
        sender_contains: Optional[str] = None,
        folder: str = "Inbox",
        max_results: int = 50,
    ) -> List[Dict[str, Any]]:
        if not self._namespace:
            self.connect()

        inbox = self._namespace.GetDefaultFolder(6)  # 6 = Inbox
        if folder != "Inbox":
            inbox = inbox.Folders[folder]

        messages = inbox.Items
        messages.Sort("[ReceivedTime]", True)

        results = []
        for msg in messages:
            if len(results) >= max_results:
                break
            if subject_contains and subject_contains.lower() not in str(msg.Subject).lower():
                continue
            if sender_contains and sender_contains.lower() not in str(msg.SenderName).lower():
                continue
            results.append({
                "subject": msg.Subject,
                "sender": msg.SenderName,
                "received": str(msg.ReceivedTime),
                "body": msg.Body[:500] if msg.Body else "",
                "has_attachments": msg.Attachments.Count > 0,
            })
        return results

    def send_email(
        self,
        to: str,
        subject: str,
        body: str,
        attachment_paths: Optional[List[str]] = None,
    ):
        if not self._outlook:
            self.connect()

        mail = self._outlook.CreateItem(0)
        mail.To = to
        mail.Subject = subject
        mail.Body = body

        if attachment_paths:
            for path in attachment_paths:
                mail.Attachments.Add(path)

        mail.Send()
        self.logger.info(f"Email sent to {to}: {subject}")

    def close(self):
        if self._outlook:
            self._outlook = None
            self._namespace = None


class WordHandler:
    def __init__(self, logger: Optional[HarnessLogger] = None):
        self.logger = logger or HarnessLogger("word")
        self._doc = None

    def read_docx(self, path: str) -> str:
        try:
            from docx import Document
            self._doc = Document(path)
            text = "\n".join(p.text for p in self._doc.paragraphs)
            self.logger.info(f"Read Word document: {path}")
            return text
        except ImportError:
            self.logger.error("python-docx not installed. Run: pip install python-docx")
            raise

    def find_in_docx(self, path: str, search_text: str) -> List[Dict[str, Any]]:
        from docx import Document
        doc = Document(path)
        matches = []
        for i, p in enumerate(doc.paragraphs):
            if search_text.lower() in p.text.lower():
                matches.append({"paragraph_index": i, "text": p.text[:200], "style": str(p.style.name)})
        for i, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if search_text.lower() in cell.text.lower():
                        matches.append({"table_index": i, "row": r_idx, "column": c_idx, "text": cell.text[:200]})
        return matches

    def extract_tables(self, path: str) -> List[List[List[str]]]:
        from docx import Document
        doc = Document(path)
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)
        self.logger.info(f"Extracted {len(tables)} tables from {path}")
        return tables

    def create_docx(self, output_path: str, content_paragraphs: List[str]):
        try:
            from docx import Document
            doc = Document()
            for text in content_paragraphs:
                doc.add_paragraph(text)
            doc.save(output_path)
            self.logger.info(f"Created Word document: {output_path}")
        except ImportError:
            self.logger.error("python-docx not installed")
            raise


class PDFHandler:
    def __init__(self, logger: Optional[HarnessLogger] = None):
        self.logger = logger or HarnessLogger("pdf")

    def read_pdf(self, path: str) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            self.logger.info(f"Read PDF: {path} ({len(reader.pages)} pages)")
            return text
        except ImportError:
            self.logger.error("pypdf not installed. Run: pip install pypdf")
            raise

    def extract_form_fields(self, path: str) -> Dict[str, Any]:
        from pypdf import PdfReader
        reader = PdfReader(path)
        fields = reader.get_fields()
        return {k: v.get("/V", "") if v else "" for k, v in (fields or {}).items()}

    def get_page_count(self, path: str) -> int:
        from pypdf import PdfReader
        reader = PdfReader(path)
        return len(reader.pages)

    def merge_pdfs(self, input_paths: List[str], output_path: str):
        from pypdf import PdfWriter, PdfReader
        writer = PdfWriter()
        for path in input_paths:
            reader = PdfReader(path)
            for page in reader.pages:
                writer.add_page(page)
        with open(output_path, "wb") as f:
            writer.write(f)
        self.logger.info(f"Merged {len(input_paths)} PDFs → {output_path}")
